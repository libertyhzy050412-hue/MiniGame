#!/usr/bin/env python3
"""Convert 游戏说明书.md to a professionally formatted .docx file."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── Style definitions ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
if rFonts is None:
    rPr = style.element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="微软雅黑"/>')
    rPr.append(rFonts)

# Heading styles
for i, (size, color_hex) in enumerate([(26, '1a1a1a'), (18, '333333'), (15, '555555'), (13, '666666')], 1):
    h = doc.styles[f'Heading {i}']
    h.font.name = '微软雅黑'
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = RGBColor.from_string(color_hex)
    h.paragraph_format.space_before = Pt(18 if i <= 2 else 14)
    h.paragraph_format.space_after = Pt(10 if i <= 2 else 8)
    rPr = h.element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="微软雅黑"/>')
    rPr.append(rFonts)

# ── Helper functions ──
def add_styled_paragraph(text, bold=False, size=11, color=None, alignment=None, spacing_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(spacing_after)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string('2d2d2d')
    # Grey background via shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="f5f2eb" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(1.2 + level * 1.0)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = '微软雅黑'
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string('ffffff')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dark header background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4a3728" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = '微软雅黑'
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Alternating row colors
            if r % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="faf7f0" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacing after table
    return table

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="c8b896"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_styled_paragraph('锈甲长歌', bold=True, size=38, color='3d2613',
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=8)
add_styled_paragraph('Armor Stained With Sorrow', bold=False, size=16, color='8c6b4a',
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=30)
add_divider()
add_styled_paragraph('完整操作指南', bold=False, size=14, color='6b5030',
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=6)
add_styled_paragraph('工程结构 · 使用方式 · 基本操作 · 全流程攻略', bold=False, size=11, color='9e8b72',
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=6)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPTER 1: 工程文件结构
# ══════════════════════════════════════════════════════════════
doc.add_heading('一、工程文件结构', level=1)

add_styled_paragraph('项目根目录结构如下：', size=10.5, color='555555')

structure = """project/
├── 游戏说明书.md                 ← 本文件
├── analysis/                     # 设计文档与需求分析
│   ├── development-plan.md       # v1 开发计划
│   ├── extracted/                # 四关文字剧本 (.txt)
│   └── pdf_pages/                # 纸面原型截图
├── Arts/                         # 原始美术资源 (.ai/.png/.jpg)
│   ├── 第一关 (1)/, 第一关补充/
│   ├── 第二关部分/
│   ├── 第三关部分/, 第三关补充/
│   └── 第四关/
├── game/                         # ★ 游戏工程主体
│   ├── index.html                # 入口 HTML
│   ├── package.json              # 项目配置 / 依赖 / 打包
│   ├── tsconfig.json             # TypeScript 配置
│   ├── vite.config.ts            # Vite 构建配置
│   ├── src/                      # TypeScript 源码
│   │   ├── main.ts               # 游戏主循环 (~2150行)
│   │   ├── narrativeData.ts      # 场景/对话/笔记数据
│   │   ├── generatedSceneArt.ts  # SVG 程序化场景图
│   │   ├── storage.ts            # 存档读写
│   │   ├── validators.ts         # 谜题校验逻辑
│   │   ├── hotspotLayout.ts      # 热点标签避让引擎
│   │   ├── style.css             # 全局样式表
│   │   └── *.test.ts             # Vitest 测试
│   ├── electron/                 # Electron 桌面壳
│   │   ├── main.cjs              # 主进程
│   │   └── preload.cjs           # 预加载脚本
│   ├── public/art/               # 运行时美术资源
│   │   ├── act1/, act2/, act3/   # 关卡场景图
│   │   └── ui/                   # UI 元件
│   ├── renderer/                 # Vite 构建产物（自动生成）
│   └── release/                  # ★ Electron 打包产物
│       ├── 锈甲长歌 Setup 0.0.0.exe    # 安装包
│       └── win-unpacked/                # 免安装版
│           └── 锈甲长歌.exe             # 双击即玩
└── tools/                        # 辅助脚本"""

add_code_block(structure)

doc.add_heading('技术栈', level=2)
add_table(
    ['层级', '技术方案'],
    [
        ['前端构建', 'Vite + TypeScript'],
        ['游戏渲染', '纯 DOM 操作（innerHTML 模板）'],
        ['状态管理', '单一 SaveState 对象（flags + currentScene）'],
        ['持久化存储', 'localStorage (键: war-and-peace-save-v1)'],
        ['桌面壳', 'Electron + electron-builder'],
        ['打包格式', 'Windows NSIS 安装包 / 免安装目录'],
        ['测试框架', 'Vitest (8 个测试, 2 组用例)'],
    ],
    col_widths=[3.5, 11.5]
)

# ══════════════════════════════════════════════════════════════
# CHAPTER 2: 文件使用方式
# ══════════════════════════════════════════════════════════════
doc.add_heading('二、文件使用方式', level=1)

doc.add_heading('2.1 玩家使用（运行游戏）', level=2)

add_styled_paragraph('免安装版（推荐快速体验）', bold=True, size=11)
add_code_block('release/win-unpacked/锈甲长歌.exe')
add_bullet('双击即可运行，无需任何安装')
add_bullet('整个 win-unpacked 文件夹可拷贝至 U 盘或任意目录')

add_styled_paragraph('安装版', bold=True, size=11)
add_code_block('release/锈甲长歌 Setup 0.0.0.exe')
add_bullet('双击运行安装向导，支持自定义安装路径')
add_bullet('自动创建桌面快捷方式和开始菜单项')
add_bullet('可通过 Windows「设置 → 应用」卸载')

doc.add_heading('2.2 开发者使用', level=2)

add_styled_paragraph('环境要求：Node.js ≥ 18，npm ≥ 9', size=10, color='888888')

add_table(
    ['命令', '作用'],
    [
        ['npm --prefix game install', '安装依赖'],
        ['npm --prefix game run dev', '开发模式（Vite + Electron 热更新）'],
        ['npm --prefix game run dev:web', '仅启动 Web 端 (http://127.0.0.1:5173)'],
        ['npm --prefix game run typecheck', 'TypeScript 类型检查'],
        ['npm --prefix game run test:run', '运行 Vitest 测试'],
        ['npm --prefix game run build:web', '生产构建（仅 Web）→ renderer/'],
        ['npm --prefix game run build', '完整打包 → release/*.exe（5-10分钟）'],
    ],
    col_widths=[6.5, 8.5]
)

# ══════════════════════════════════════════════════════════════
# CHAPTER 3: 游戏基本操作
# ══════════════════════════════════════════════════════════════
doc.add_heading('三、游戏基本操作', level=1)

doc.add_heading('3.1 通用操作', level=2)
add_table(
    ['操作', '方式'],
    [
        ['交互 / 调查', '点击场景中的圆形光点（热点）'],
        ['推进对话', '点击对话框任意位置'],
        ['选择选项', '点击谜题面板中的按钮'],
        ['关闭面板', '点击右上角 × 按钮'],
        ['查看线索簿', '点击右下角「线索簿」按钮'],
        ['返回主菜单', '点击右下角「主菜单」按钮'],
        ['重新开始', '点击右下角「重置」或主菜单「从头开始」'],
    ],
    col_widths=[4, 11]
)

doc.add_heading('3.2 各谜题专属操作', level=2)
add_table(
    ['谜题', '操作方式'],
    [
        ['棋局残局', '按顺序点击三个落子选项'],
        ['圣器对应 / 法阵摆放', '为每项点击选择对应答案'],
        ['舞者谜语', '从三个选项中点击选择答案'],
        ['生火', '点击火石 / 火绒 / 鼓风按钮，调节温度和风力至绿色区间'],
        ['家书匹配', '按住家书卡片拖拽到右侧对应士兵栏位'],
        ['伤兵登记', '点击选择 3 名真正受伤的士兵'],
        ['搜救村民', 'W / A / S / D 键移动，推开箱子走到村民位置'],
        ['赶羊', '按住小羊拖拽到画面中央的绿色圆形靶心，倒计时归零时全部在圈内即胜'],
    ],
    col_widths=[4, 11]
)

doc.add_heading('3.3 界面布局', level=2)
add_code_block("""┌──────────────────────────────────────────┐
│  左上角     当前章节名 + 场景标题        │
│                                          │
│          [ 游戏场景画面 ]                │
│     ● 热点（圆形光点 + 标签）            │
│                                          │
│  左下角     当前目标提示                 │
│  右下角     线索簿 | 主菜单 | 重置       │
│                                          │
│  [ 对话框区域 — 点击推进对话 ]           │
└──────────────────────────────────────────┘""")

# ══════════════════════════════════════════════════════════════
# CHAPTER 4: 全流程攻略
# ══════════════════════════════════════════════════════════════
doc.add_heading('四、主线全流程攻略', level=1)

# --- Act 1 ---
doc.add_heading('第一关 · 书房的秘密', level=2)
add_styled_paragraph('剧情概要：农家青年受领主举荐前往都城王宫任书记官，通过调查文书、账目和密信，揭开国王发动战争的财政困境与外部勒索。', size=10, color='777777')

scenes_act1 = [
    ('场景 1：农舍（送别）', [
        '自动播放父母送别对话',
        '对话结束后点击「启程」热点前往王宫',
    ]),
    ('场景 2：王宫门口', [
        '① 点击左侧小丑对话，了解布告栏和书房信息',
        '② 点击右侧布告栏阅读三份布告（税赋 / 异端清查 / 征兵招募）',
        '③ 阅读完毕后点击中央守卫应聘书记官',
    ]),
    ('场景 3：谒见国王', [
        '自动播放国王任命对话',
        '对话结束后点击左下「去二楼」',
    ]),
    ('场景 4：二楼走廊', [
        '（可选）依次点击四幅画像，查看历代国王自述',
        '点击右侧「书房门」进入书房',
    ]),
    ('场景 5：书房调查 ★ 核心', [
        '① 征兵起草书 — 桌面中央，国王征召 16-60 岁男子',
        '② 财政报告 — 桌面右侧，国库赤字 16080 英镑',
        '③ 搜寻桌底 — 先点搜寻，再点出现的密信（敌国勒索信）',
        '④ 残局棋盘 — 解开棋局获得兵力报告（仅 1870 正规军）',
        '（可选）书架上历书 — 查看四代国王变迁',
        '全部调查完成后点右上角「面见国王」，自动跳转',
    ]),
]
for title, items in scenes_act1:
    add_styled_paragraph(title, bold=True, size=11.5)
    for item in items:
        add_bullet(item)

doc.add_heading('谜题 A：残局棋局', level=3)
add_table(
    ['步骤', '选择'],
    [['第一步', '白后逼近王侧'], ['第二步', '白象封住退路'], ['第三步', '白后完成将杀']],
    col_widths=[3, 12]
)

# --- Act 2 ---
doc.add_heading('第二关 · 圣器的试炼', level=2)
add_styled_paragraph('剧情概要：主角前往教堂为出征祷告，却发现宗教仪式将战争包装成神圣使命。通过经文研读和圣器试炼，逐渐看清战争正当性的虚伪一面。', size=10, color='777777')

add_styled_paragraph('场景 6：教堂与高塔', bold=True, size=11.5)
add_bullet('① 点击中央「祷告」，听取教皇战前祷告')
add_bullet('② 祷告后左侧出现「藏经室」，点击阅读经文')
add_bullet('③ 阅读后出现「圣器试炼」，点击进入')
add_bullet('④ 试炼完成后右侧出现「舞者与高塔」，点击进入谜语')
add_bullet('⑤ 答完三道谜语后，点击「奔赴军营」')

doc.add_heading('谜题 B：圣器与教义对应', level=3)
add_table(
    ['经文摘要', '正确圣器'],
    [
        ['"不可跪拜那些像，也不可事奉它们……"', '金牛形巴力像'],
        ['"你要在米甸人身上报仇……"', '米甸血矛'],
        ['"三百人呐喊，使他们逃跑……"', '基甸羊角号'],
        ['"凡能见火的，必过火洁净……"', '祭司纯金胸牌'],
    ],
    col_widths=[7.5, 7.5]
)

doc.add_heading('谜题 C：法阵方位摆放', level=3)
add_table(
    ['方位', '正确神兽'],
    [['正北', '羔羊'], ['正东', '蛇'], ['正南', '鹰'], ['正西', '狮']],
    col_widths=[7.5, 7.5]
)

doc.add_heading('谜题 D：舞者三道谜语', level=3)
add_table(
    ['问题', '正确答案'],
    [
        ['是什么艳色诱人眼目，摘取便动真心？', '夏娃的苹果'],
        ['是什么利锋可破沉淤，矫健而辟新章？', '沾血的长剑'],
        ['是什么光彩引人心动，世人皆愿追逐？', '闪耀的金币'],
    ],
    col_widths=[7.5, 7.5]
)

# --- Act 3 ---
doc.add_heading('第三关 · 战争的代价', level=2)
add_styled_paragraph('剧情概要：主角抵达前线，在生火、医帐登记、村庄救援和战场中目睹战争的真实代价——补给匮乏、平民受害、宣传与现实的巨大落差。', size=10, color='777777')

add_styled_paragraph('场景 7：寒夜军营', bold=True, size=11.5)
add_bullet('① 点击「木堆」进入生火小游戏')
add_bullet('② 生火成功后与「骑士」对话接受命令')
add_bullet('③ 点击右下「去医帐」')

doc.add_heading('谜题 E：生火', level=3)
add_styled_paragraph('保持火芯温度 38-76 和鼓风力度 34-72 在绿色区间内：', size=10.5)
add_table(
    ['按钮', '温度变化', '风力变化', '进度变化'],
    [
        ['🔥 敲击火石', '+18', '-8', '—'],
        ['🌿 添引火绒', '+10', '-4', '+7~18'],
        ['💨 缓缓鼓风', '+2~7', '+16', '+4~14'],
    ],
    col_widths=[4, 3.5, 3.5, 4]
)
add_styled_paragraph('提示：先用火石升温，再用鼓风调风力，最后反复添火绒。温度和风力都在绿色区间时进度上涨最快。', size=10, color='888888')

add_styled_paragraph('场景 8：医帐', bold=True, size=11.5)
add_bullet('① 点击左侧「家书」进入匹配')
add_bullet('② 点击右侧「伤兵登记」找出 3 名真正受伤的士兵')
add_bullet('③ 点击右下「列阵出发」自动跳转')

doc.add_heading('谜题 F：家书匹配', level=3)
add_table(
    ['遗体特征', '对应家书', '关键描述'],
    [
        ['手上戴戒指，手臂带伤', '家书一', '左肩受创，药很少，只草草包扎'],
        ['头发花白，旧伤累累', '家书二', '年岁已老，求卸甲归田'],
        ['个头更矮，像刚离家不久', '家书三', '晚餐黑麦面包和土豆，想念家里的床'],
    ],
    col_widths=[4.5, 3, 7.5]
)

doc.add_heading('谜题 G：伤兵登记', level=3)
add_styled_paragraph('选择以下 3 名真正受伤的士兵：', size=10.5)
add_table(
    ['选择', '士兵', '伤情'],
    [
        ['✅', '弩手', '左肩渗血绷带，披风半垂'],
        ['✅', '长枪兵', '右腿木板固定，只能借枪杆撑地'],
        ['✅', '盾兵', '腹侧临时压着布团，呼吸急促'],
        ['❌', '辎重兵', '只是满脸烟灰，能自己站直说话'],
        ['❌', '传令兵', '裹着毯子发抖，但没有新伤口'],
        ['❌', '伙夫', '手背起泡，但还能拎着水桶走动'],
        ['❌', '老兵', '旧伤很多，但此刻没有新的出血'],
    ],
    col_widths=[1.5, 3, 10.5]
)

add_styled_paragraph('场景 9：列阵出发', bold=True, size=11.5)
add_bullet('点击「继续前进」')

add_styled_paragraph('场景 10：火中的村庄', bold=True, size=11.5)
add_bullet('① 点击左上「藏身点」进入搜救（WASD 推箱子，救出 3 名村民）')
add_bullet('② 依次点击「战时税令」和「军饷文书」阅读')
add_bullet('③ 点击右上「战场」')

doc.add_heading('谜题 H：搜救村民（WASD 推箱子）', level=3)
add_bullet('使用 W/A/S/D 键移动角色，推开箱子，走到村民（金色小人）所在格子即救出')
add_bullet('3 名村民分别位于右上角、右下角、左下角')
add_bullet('箱子只能推动一格，需前方为空地且无其他箱子阻挡')

add_styled_paragraph('场景 11：失序战场', bold=True, size=11.5)
add_bullet('自动播放战场昏迷过场')

# --- Act 4 ---
doc.add_heading('第四关 · 战争的终结', level=2)
add_styled_paragraph('剧情概要：主角负伤后被村民救回农舍。在停战前夕的村庄里，平民对战争的解释与王宫、教堂完全不同——和平不是宣传口号，而是能否好好过日子。', size=10, color='777777')

add_styled_paragraph('场景 12：苏醒', bold=True, size=11.5)
add_bullet('自动播放农妇对话，点击「去村里」')

add_styled_paragraph('场景 13：停战前夕的村庄', bold=True, size=11.5)
add_bullet('① 点击左侧「村民」听三种不同的战争立场')
add_bullet('② 点击右侧「羊圈」进入赶羊小游戏')
add_bullet('③ 点击中央「孩子们」听平民版谜语答案')
add_bullet('④ 点击右下「送信者」听取停战消息')

doc.add_heading('谜题 I：赶羊', level=3)
add_bullet('羊从牧场中心向四个方向快速跑开')
add_bullet('按住并拖动小羊，把它们拖回画面中央的绿色圆形靶心')
add_bullet('羊被拖入圈内后困住 1 秒（变为半透明绿色），随后重新向外跑')
add_bullet('30 秒倒计时归零时，所有 4 只羊都在圆形靶心内即为胜利')
add_bullet('一旦有羊跑出画面边界，游戏立即失败')

add_styled_paragraph('场景 14：战争之后（结局）', bold=True, size=11.5)
add_bullet('阅读结局总结，可点击返回主菜单或重新开始')

# ══════════════════════════════════════════════════════════════
# CHAPTER 5: 存档
# ══════════════════════════════════════════════════════════════
doc.add_heading('五、存档与进度', level=1)

add_table(
    ['功能', '说明'],
    [
        ['自动存档', '使用浏览器 localStorage，每完成对话 / 解谜 / 切场景即时保存'],
        ['存档键名', 'war-and-peace-save-v1'],
        ['继续旅程', '从上次退出位置继续（保留所有已获得旗标）'],
        ['从头开始', '清空存档，回到农舍开场'],
        ['线索簿', '右下角按钮，左侧列表显示已读条目，点击查看完整内容'],
    ],
    col_widths=[4, 11]
)

# ══════════════════════════════════════════════════════════════
# CHAPTER 6: FAQ
# ══════════════════════════════════════════════════════════════
doc.add_heading('六、常见问题', level=1)

doc.add_heading('运行问题', level=2)
faq_run = [
    ('Q：双击 exe 没反应？', '检查 win-unpacked 文件夹是否完整（约 200MB）。杀毒软件可能拦截，尝试添加信任。'),
    ('Q：窗口显示不全？', '游戏设计分辨率为 16:9，最低 1280×720，建议 1920×1080。开发模式按 F11 全屏。'),
    ('Q：如何卸载？', '安装版：Windows「设置 → 应用 → 锈甲长歌 → 卸载」。免安装版：直接删除文件夹。'),
]
for q, a in faq_run:
    add_styled_paragraph(q, bold=True, size=11, spacing_after=2)
    add_styled_paragraph(a, size=10.5, color='555555', spacing_after=10)

doc.add_heading('游戏问题', level=2)
faq_game = [
    ('Q：生火进度条不动？', '温度和风力必须同时在绿色区间内，进度才会稳定上涨。单独操作一项很难成功。'),
    ('Q：推箱子卡住了？', '箱子只能推动一格，前方必须是空地且无其他箱子阻挡。如卡死，点击「重新搜路」重来。'),
    ('Q：赶羊太难？', '优先拖住跑得最快、最接近画面边缘的羊。拖回绿色圆形靶心后羊会困住 1 秒再跑。倒计时最后几秒要确保所有羊都在圈内。'),
    ('Q：对话结束后不知去哪？', '观察场景中发光的圆形热点。线索簿会显示当前场景标题帮助你判断进度。'),
]
for q, a in faq_game:
    add_styled_paragraph(q, bold=True, size=11, spacing_after=2)
    add_styled_paragraph(a, size=10.5, color='555555', spacing_after=10)

doc.add_heading('开发问题', level=2)
faq_dev = [
    ('Q：如何新增场景？', '在 narrativeData.ts 中添加 SceneDefinition，在 main.ts 的 handleAction 中注册对应 action。'),
    ('Q：如何新增谜题？', '在 main.ts 中定义 OverlayState 接口、编写 render*Puzzle 函数、在 handleCommand 中添加命令处理。'),
    ('Q：如何替换美术资源？', '将 PNG 放入 public/art/ 对应目录，修改 narrativeData.ts 中场景的 background 路径。缺失场景可在 generatedSceneArt.ts 中用 SVG 补画。'),
]
for q, a in faq_dev:
    add_styled_paragraph(q, bold=True, size=11, spacing_after=2)
    add_styled_paragraph(a, size=10.5, color='555555', spacing_after=10)

# ── Save ──
output_path = '游戏说明书.docx'
doc.save(output_path)
print(f'Done → {output_path}')
